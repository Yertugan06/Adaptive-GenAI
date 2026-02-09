import fitz
from docx import Document
from typing import List, Optional
from backend.schemas.nosql.document_chunk import DocumentChunk
from backend.services.bi_encoder import count_tokens, create_embedding

class DocumentProcessor:

    MAX_MODEL_TOKENS = 512
    CHUNK_SIZE_TOKENS = 400  
    CHUNK_OVERLAP_TOKENS = 50
    
    @staticmethod
    def validate_token_count(text: str) -> bool:
        """Check if text exceeds model token limit."""
        return count_tokens(text) <= DocumentProcessor.MAX_MODEL_TOKENS

    @staticmethod
    def truncate_to_token_limit(text: str, max_tokens: int = None) -> str: # type: ignore
        """Safely truncate text to not exceed token limit."""
        if max_tokens is None:
            max_tokens = DocumentProcessor.MAX_MODEL_TOKENS
        
        tokens = text.split()
        if not tokens:
            return text
            
        truncated_text = ""
        for word in tokens:
            test_text = f"{truncated_text} {word}".strip()
            if count_tokens(test_text) <= max_tokens:
                truncated_text = test_text
            else:
                break
        
        return truncated_text

    @staticmethod
    def get_chunks_from_text(text: str, parent_id: str, company_id: int, 
                             page: Optional[int] = None) -> List[DocumentChunk]:

        text = text.strip()
        if not text:
            return []
        
        if count_tokens(text) <= DocumentProcessor.CHUNK_SIZE_TOKENS:
            if DocumentProcessor.validate_token_count(text):
                return [DocumentChunk(
                    parent_doc_id=parent_id,
                    company_id=company_id,
                    chunk_index=0,
                    content=text,
                    embedding=create_embedding(text),
                    page_number=page
                )]
            else:
                safe_text = DocumentProcessor.truncate_to_token_limit(
                    text, DocumentProcessor.CHUNK_SIZE_TOKENS
                )
                return [DocumentChunk(
                    parent_doc_id=parent_id,
                    company_id=company_id,
                    chunk_index=0,
                    content=safe_text,
                    embedding=create_embedding(safe_text),
                    page_number=page
                )]
        
        sentences = []
        current_sentence = []
        

        for char in text:
            current_sentence.append(char)
            if char in '.!?。！？':
                sentence = ''.join(current_sentence).strip()
                if sentence:
                    sentences.append(sentence)
                current_sentence = []
        
        if current_sentence:
            sentence = ''.join(current_sentence).strip()
            if sentence:
                sentences.append(sentence)
        
        chunks = []
        current_chunk_sentences = []
        current_chunk_text = ""
        chunk_idx = 0
        
        for sentence in sentences:
            potential_chunk = f"{current_chunk_text} {sentence}".strip() if current_chunk_text else sentence
            
            if count_tokens(potential_chunk) > DocumentProcessor.CHUNK_SIZE_TOKENS:
                if current_chunk_text:

                    if not DocumentProcessor.validate_token_count(current_chunk_text):
                        current_chunk_text = DocumentProcessor.truncate_to_token_limit(
                            current_chunk_text, DocumentProcessor.CHUNK_SIZE_TOKENS
                        )
                    
                    chunks.append(DocumentChunk(
                        parent_doc_id=parent_id,
                        company_id=company_id,
                        chunk_index=chunk_idx,
                        content=current_chunk_text,
                        embedding=create_embedding(current_chunk_text),
                        page_number=page
                    ))
                    chunk_idx += 1
                    
                    if DocumentProcessor.CHUNK_OVERLAP_TOKENS > 0:
                        overlap_text = DocumentProcessor.create_overlap_text(
                            current_chunk_text, 
                            DocumentProcessor.CHUNK_OVERLAP_TOKENS
                        )
                        current_chunk_text = overlap_text
                        current_chunk_sentences = [s for s in current_chunk_sentences 
                                                 if s in overlap_text]
                    else:
                        current_chunk_text = ""
                        current_chunk_sentences = []
                

                if count_tokens(sentence) > DocumentProcessor.CHUNK_SIZE_TOKENS:

                    sentence_chunks = DocumentProcessor.split_long_sentence(sentence)
                    for sub_sentence in sentence_chunks:

                        if count_tokens(sub_sentence) > DocumentProcessor.CHUNK_SIZE_TOKENS:
                            sub_chunks = DocumentProcessor.get_chunks_from_text(
                                sub_sentence, parent_id, company_id, page
                            )

                            for sub_chunk in sub_chunks:
                                sub_chunk.chunk_index = chunk_idx
                                chunks.append(sub_chunk)
                                chunk_idx += 1
                        else:
                            current_chunk_text = sub_sentence
                            current_chunk_sentences = [sub_sentence]
                else:
                    current_chunk_text = sentence
                    current_chunk_sentences = [sentence]
            else:

                if current_chunk_text:
                    current_chunk_text = f"{current_chunk_text} {sentence}"
                else:
                    current_chunk_text = sentence
                current_chunk_sentences.append(sentence)
        

        if current_chunk_text:
            if not DocumentProcessor.validate_token_count(current_chunk_text):
                current_chunk_text = DocumentProcessor.truncate_to_token_limit(
                    current_chunk_text, DocumentProcessor.CHUNK_SIZE_TOKENS
                )
            
            chunks.append(DocumentChunk(
                parent_doc_id=parent_id,
                company_id=company_id,
                chunk_index=chunk_idx,
                content=current_chunk_text,
                embedding=create_embedding(current_chunk_text),
                page_number=page
            ))
        
        return chunks
    
    @staticmethod
    def create_overlap_text(text: str, overlap_tokens: int) -> str:
        """Create overlap text from the end of a chunk."""
        if overlap_tokens <= 0:
            return ""
        

        sentences = [s.strip() for s in text.split('.') if s.strip()]
        if not sentences:
            return ""
        

        overlap_sentences = []
        current_tokens = 0
        
        for sentence in reversed(sentences):
            sentence_tokens = count_tokens(sentence)
            if current_tokens + sentence_tokens > overlap_tokens:
                break
            overlap_sentences.insert(0, sentence)
            current_tokens += sentence_tokens
        
        return '. '.join(overlap_sentences) + '.'
    
    @staticmethod
    def split_long_sentence(sentence: str) -> List[str]:
        """Split a very long sentence into smaller parts."""
        words = sentence.split()
        if len(words) <= 10:  
            return [sentence]
        

        split_points = []
        for i, word in enumerate(words):
            if word.lower() in ['and', 'but', 'or', 'however', 'therefore', 'moreover']:
                split_points.append(i)
        
        if split_points:
            parts = []
            start = 0
            for point in split_points:
                part = ' '.join(words[start:point+1])
                if part.strip():
                    parts.append(part)
                start = point + 1
            
            # Add remaining part
            if start < len(words):
                part = ' '.join(words[start:])
                if part.strip():
                    parts.append(part)
            
            return parts
        
        part_size = max(10, len(words) // 3) 
        parts = []
        for i in range(0, len(words), part_size):
            part = ' '.join(words[i:i+part_size])
            if part.strip():
                parts.append(part)
        
        return parts

    @classmethod
    def process_pdf(cls, file_path: str, parent_id: str, company_id: int) -> List[DocumentChunk]:
        """Process PDF file with page-level chunking."""
        all_chunks = []
        try:
            doc = fitz.open(file_path)
            for page_num, page in enumerate(doc, start=1): # type: ignore
                text = page.get_text()
                if text.strip():
                    page_chunks = cls.get_chunks_from_text(
                        text, parent_id, company_id, page=page_num
                    )
                    all_chunks.extend(page_chunks)
        except Exception as e:
            print(f"Error processing PDF {file_path}: {e}")
            raise
        return all_chunks

    @classmethod
    def process_docx(cls, file_path: str, parent_id: str, company_id: int) -> List[DocumentChunk]:
        """Process DOCX file with paragraph-aware chunking."""
        try:
            doc = Document(file_path)
            
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            paragraphs.append(text)
            
            full_text = "\n\n".join(paragraphs)
            
            return cls.get_chunks_from_text(full_text, parent_id, company_id)
        except Exception as e:
            print(f"Error processing DOCX {file_path}: {e}")
            raise